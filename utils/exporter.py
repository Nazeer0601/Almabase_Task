from docx import Document

def export_answers(results, filename="output.docx"):
    doc = Document()
    doc.add_heading("Completed Questionnaire", level=1)

    for r in results:
        doc.add_heading(r["question"], level=2)

        doc.add_paragraph("Answer:")
        doc.add_paragraph(r["answer"])

        doc.add_paragraph("Citation:")
        doc.add_paragraph(r["citation"])

        doc.add_paragraph("Confidence Score:")
        doc.add_paragraph(str(r["confidence"]))

        doc.add_page_break()

    doc.save(filename)